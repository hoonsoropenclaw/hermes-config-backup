#!/usr/bin/env python3
"""
generate_contract_substitute.py — Substitute teacher employment contract using python-docx (python3.12)
Usage: python3.12 generate_contract_substitute.py <candidate_name> <subject> <hourly_rate> <substitution_reason> <period> <school_name> [output_path]
"""
import sys

def main():
    if len(sys.argv) < 7:
        print("Usage: python3.12 generate_contract_substitute.py <candidate> <subject> <hourly_rate> <sub_reason> <period> <school> [output]", file=sys.stderr)
        sys.exit(1)
    candidate, subject, hourly_rate, sub_reason, period, school = sys.argv[1:7]
    output = sys.argv[7] if len(sys.argv) > 7 else f"/tmp/contract_{candidate.replace(' ', '_')}.docx"

    import subprocess
    result = subprocess.run(
        ["python3.12", "-c", """
import sys
sys.path.insert(0, '/usr/lib/python3.12/site-packages')
from docx import Document
candidate, subject, hourly_rate, sub_reason, period, school, output = sys.argv[1:8]
doc = Document()
doc.add_heading('代理教師聘用合約', 0)
doc.add_heading('Employment Contract for Substitute Teacher', level=2)
doc.add_paragraph(f'甲方（學校）: {school}')
doc.add_paragraph(f'乙方（教師）: {candidate}')
doc.add_heading('第一條 聘用期間', level=1)
doc.add_paragraph(f'甲方聘乙方為代理教師，代理期間：{period}')
doc.add_paragraph(f'代理原因：{sub_reason}')
doc.add_heading('第二條 授課科目', level=1)
doc.add_paragraph(f'乙方應授科目：{subject}')
doc.add_heading('第三條 鐘點費', level=1)
doc.add_paragraph(f'鐘點費率：每節 {hourly_rate} 元（含勞健保）')
doc.add_paragraph('計算方式：實際授課節數 × 鐘點費率')
doc.add_heading('第四條 權利義務', level=1)
doc.add_paragraph('乙方應遵守學校規章、履行教師職責、參加校內會議及研習活動。')
doc.add_paragraph('')
doc.add_paragraph('甲方簽章：________________     日期：____________')
doc.add_paragraph('乙方簽章：________________     日期：____________')
doc.save(output)
print(f'DONE:{output}')
""", candidate, subject, hourly_rate, sub_reason, period, school, output],
        capture_output=True, text=True
    )
    if result.returncode == 0 and "DONE:" in result.stdout:
        print(f"✅ Generated: {output}")
    else:
        print(f"❌ Generation failed: {result.stderr}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
