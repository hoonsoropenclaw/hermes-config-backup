#!/usr/bin/env python3
"""
generate_offer_letter.py — Generate HR offer letter DOCX using python-docx (python3.12)
Usage: python3.12 generate_offer_letter.py <candidate_name> <position> <salary> <start_date> <school_name> [output_path]
"""
import sys, os
from datetime import date

def main():
    if len(sys.argv) < 6:
        print("Usage: python3.12 generate_offer_letter.py <candidate> <position> <salary> <start_date> <school> [output]", file=sys.stderr)
        sys.exit(1)
    candidate, position, salary, start_date, school = sys.argv[1:6]
    output = sys.argv[6] if len(sys.argv) > 6 else f"/tmp/offer_letter_{candidate.replace(' ', '_')}.docx"

    # Dynamically import from python3.12 site-packages
    import subprocess
    result = subprocess.run(
        ["python3.12", "-c", """
import sys
sys.path.insert(0, '/usr/lib/python3.12/site-packages')
from docx import Document
from datetime import date
import sys
candidate, position, salary, start_date, school, output = sys.argv[1:7]
doc = Document()
doc.add_heading('錄取通知書 Offer Letter', 0)
doc.add_paragraph(f'日期/Date: {date.today().strftime("%Y年%m月%d日")}')
doc.add_paragraph(f'親愛的 {candidate}：')
doc.add_paragraph(f'恭喜您通過「{position}」一職的面試，我們誠摯邀請您加入 {school}。')
doc.add_heading('聘用條件 Employment Terms', level=1)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
data = [
    ('職位 / Position', position),
    ('到職日 / Start Date', start_date),
    ('月薪 / Monthly Salary', salary),
    ('聘用期間 / Employment Period', f'{start_date} ~ {int(start_date[:4])+1}-07-31'),
    ('適用法規 / Applicable Law', '教師法、教育部代理教師注意事項'),
]
for i, (key, val) in enumerate(data):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = val
doc.add_paragraph('')
doc.add_paragraph('請在收到此通知後 5 個工作天內回覆是否接受此錄取。')
doc.add_paragraph('若有任何問題，請聯繫人事部門。')
doc.add_paragraph(f'\\n{school} 人事部門')
doc.save(output)
print(f'DONE:{output}')
""", candidate, position, salary, start_date, school, output],
        capture_output=True, text=True
    )
    if result.returncode == 0 and "DONE:" in result.stdout:
        print(f"✅ Generated: {output}")
    else:
        print(f"❌ Generation failed: {result.stderr}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
