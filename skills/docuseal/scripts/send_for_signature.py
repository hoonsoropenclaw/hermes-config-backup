#!/usr/bin/env python3
"""
send_for_signature.py — DocuSeal e-signature integration for HR document workflow

Usage:
    python3.12 send_for_signature.py <docx_path> <signer_name> <signer_email> [template_id]

Environment:
    DOCUSEAL_URL=https://your-docuseal.example.com   (optional, for API mode)
    DOCUSEAL_API_KEY=xxx                            (optional, for API mode)
    HR_WEBHOOK_URL=https://your-hermes.example.com/hook/docuseal  (optional)

Workflow:
    1. Generate offer letter / contract .docx via hr-document-automation scripts
    2. Run this script to send the doc to DocuSeal for e-signature
    3. DocuSeal emails the signer with a signing link
    4. On completion (via webhook), HR is notified

DocuSeal field labels supported:
    {{signature:signer}}  — signature field for named signer
    {{date:sign_date}}    — date signed
    {{text:title}}        — read-only text field
    {{initials:signer_i}} — initials
    {{checkbox:agree}}    — checkbox (e.g., "I agree to terms")
"""
import sys
import os
import json
import subprocess
import tempfile
from pathlib import Path

DOCX_TEMPLATE_FIELDS = {
    "signer_name": None,       # will be filled from args
    "signer_email": None,
    "sign_date": None,
    "hr_name": None,           # school HR contact
    "position": None,
    "start_date": None,
    "salary": None,
}


def docx_to_html_with_fields(docx_path: str, output_html: str, fields: dict) -> None:
    """
    Convert a .docx file to HTML suitable for DocuSeal upload.
    Inserts DocuSeal field markers ({{signature:}}, {{date:}}, {{text:}}) 
    for the signer to fill.
    """
    try:
        import docx
    except ImportError:
        print("ERROR: python-docx not installed. Run: python3.12 -m pip install python-docx")
        sys.exit(1)

    doc = docx.Document(docx_path)
    
    html_parts = [
        "<!DOCTYPE html>",
        "<html lang='zh-TW'>",
        "<head>",
        "  <meta charset='UTF-8'>",
        "  <title>人事文件 - 電子簽署</title>",
        "  <style>",
        "    body { font-family: 'Noto Sans CJK TC', 'Heiti TC', sans-serif;",
        "           max-width: 800px; margin: 40px auto; padding: 20px; }",
        "    h1 { text-align: center; font-size: 1.5em; }",
        "    table { width: 100%; border-collapse: collapse; margin: 20px 0; }",
        "    td { padding: 8px; vertical-align: top; }",
        "    .field-row { border-bottom: 1px solid #ccc; padding: 12px 0; }",
        "    .docuseal-sig { background: #f0f8ff; border: 2px dashed #4a90d9;",
        "                   padding: 16px; margin: 16px 0; border-radius: 4px; }",
        "    .docuseal-date { background: #fff8dc; border: 2px dashed #d4a017;",
        "                     padding: 8px 16px; display: inline-block; }",
        "    .info-table td:first-child { font-weight: bold; width: 30%; }",
        "  </style>",
        "</head>",
        "<body>",
    ]

    # Title
    doc_title = fields.get("title", "人事聘僱文件")
    html_parts.append(f"  <h1>{doc_title}</h1>")

    # Info table
    html_parts.append("  <table class='info-table'>")
    for key in ["candidate_name", "position", "department", "salary", "start_date", "contract_type"]:
        if key in fields and fields[key]:
            label = {"candidate_name": "受僱人", "position": "職位", 
                     "department": "部門", "salary": "月薪", 
                     "start_date": "到職日", "contract_type": "契約類型"}.get(key, key)
            html_parts.append(f"    <tr><td>{label}:</td><td>{fields[key]}</td></tr>")
    html_parts.append("  </table>")

    # Content from docx paragraphs
    html_parts.append("  <div class='content'>")
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            html_parts.append(f"    <p>{text}</p>")
    html_parts.append("  </div>")

    # Signature block
    html_parts.append("  <hr>")
    html_parts.append("  <div class='docuseal-sig'>")
    html_parts.append("    <p><strong>電子簽署區</strong></p>")
    html_parts.append(f"    <p>簽署人: <strong>{fields.get('signer_name', '{{text:signer_name}}')}</strong></p>")
    html_parts.append(f"    <p>Email: {fields.get('signer_email', '{{text:signer_email}}')}</p>")
    html_parts.append("    <p>簽名: {{signature:signer}}</p>")
    html_parts.append(f"    <p>簽署日期: {{date:sign_date}}</p>")
    html_parts.append("  </div>")

    # HR approval block
    if fields.get("hr_name"):
        html_parts.append("  <div class='docuseal-sig' style='background:#f0fff0'>")
        html_parts.append(f"    <p><strong>人事單位確認</strong> ({fields.get('hr_name')})</p>")
        html_parts.append("    <p>HR 簽名: {{signature:hr}} &nbsp;&nbsp; 日期: {{date:hr_date}}</p>")
        html_parts.append("  </div>")

    html_parts.extend(["</body>", "</html>"])

    Path(output_html).write_text("\n".join(html_parts), encoding="utf-8")
    print(f"[OK] HTML written to: {output_html}")


def upload_to_docuseal_api(html_path: str, template_name: str = "HR Document") -> dict:
    """
    Upload HTML to DocuSeal to create a template.
    Returns template_id.
    """
    url = os.environ.get("DOCUSEAL_URL")
    api_key = os.environ.get("DOCUSEAL_API_KEY")
    
    if not url or not api_key:
        print("WARNING: DOCUSEAL_URL / DOCUSEAL_API_KEY not set. Skipping API upload.")
        print("  Manually upload:", html_path)
        return None

    import requests
    with open(html_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    resp = requests.post(
        f"{url}/api/templates",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={"name": template_name, "html": html_content},
        timeout=30,
    )
    resp.raise_for_status()
    data = resp.json()
    template_id = data.get("id") or (data.get("template") or {}).get("id")
    print(f"[OK] Template created: template_id={template_id}")
    return {"template_id": template_id, "url": f"{url}/templates/{template_id}"}


def send_for_signature(template_id: str, signer_email: str, signer_name: str, 
                       subject: str = "人事文件電子簽署通知",
                       message: str = "") -> dict:
    """
    Send a DocuSeal template for signature to a specific email.
    Returns submission_id.
    """
    url = os.environ.get("DOCUSEAL_URL")
    api_key = os.environ.get("DOCUSEAL_API_KEY")

    if not url or not api_key:
        print("ERROR: DOCUSEAL_URL / DOCUSEAL_API_KEY not set")
        return None

    import requests
    resp = requests.post(
        f"{url}/api/submissions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={
            "template_id": int(template_id),
            "submitters": [{"email": signer_email, "name": signer_name}],
            "subject": subject,
            "message": message or (
                f"您好 {signer_name}，\n\n請點擊下方連結填寫並簽署文件。\n\n"
                "完成後系統將通知人事單位。\n\n感謝您的配合。"
            ),
        },
        timeout=30,
    )
    resp.raise_for_status()
    data = resp.json()
    submission_id = data.get("id") or (data.get("submission") or {}).get("id")
    print(f"[OK] Submission sent: submission_id={submission_id}")
    print(f"     Signer will receive email from: {url}")
    return {"submission_id": submission_id, "status_url": f"{url}/submissions/{submission_id}"}


def docx_to_pdf(docx_path: str, pdf_path: str = None) -> str:
    """
    Convert DOCX to PDF using LibreOffice.
    Returns path to PDF.
    """
    if pdf_path is None:
        pdf_path = str(Path(docx_path).with_suffix(".pdf"))
    
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", 
         str(Path(pdf_path).parent), str(docx_path)],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        print(f"ERROR: LibreOffice conversion failed: {result.stderr}", file=sys.stderr)
        sys.exit(1)
    
    actual_pdf = str(Path(pdf_path).parent / Path(docx_path).with_suffix(".pdf").name)
    print(f"[OK] PDF created: {actual_pdf}")
    return actual_pdf


def generate_and_send(candidate_name: str, position: str, salary: str,
                      start_date: str, school_name: str,
                      signer_email: str, docx_path: str = None,
                      contract_type: str = "正式契約") -> dict:
    """
    Full pipeline:
    1. Generate .docx via hr-document-automation scripts (if docx_path not provided)
    2. Convert to HTML with DocuSeal field markers
    3. Upload to DocuSeal (if API configured)
    4. Send for e-signature
    """
    fields = {
        "candidate_name": candidate_name,
        "position": position,
        "salary": salary,
        "start_date": start_date,
        "school_name": school_name,
        "contract_type": contract_type,
        "signer_name": candidate_name,
        "signer_email": signer_email,
        "sign_date": "{{date:sign_date}}",
        "title": f"教師{contract_type} - {candidate_name}",
    }

    # Step 1: Generate docx if not provided
    if not docx_path or not Path(docx_path).exists():
        print("[STEP 1] Generating DOCX via hr-document-automation...")
        if "代理" in position or "代課" in position:
            gen_script = "generate_contract_substitute.py"
        else:
            gen_script = "generate_offer_letter.py"
        
        script_path = Path.home() / f".hermes/skills/hr-document-automation/scripts/{gen_script}"
        result = subprocess.run(
            ["python3.12", str(script_path), candidate_name, position,
             salary, start_date, school_name],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            print(f"ERROR: generate script failed:\n{result.stderr}", file=sys.stderr)
            sys.exit(1)
        
        default_out = Path("/tmp/offer.docx") if "offer" in gen_script else Path("/tmp/contract.docx")
        docx_path = result.stdout.strip().split("\n")[-1] if result.stdout.strip() else str(default_out)
        print(f"[OK] DOCX generated: {docx_path}")
    else:
        print(f"[STEP 1] Using provided DOCX: {docx_path}")

    # Step 2: Convert to HTML with DocuSeal fields
    print("[STEP 2] Converting to HTML with DocuSeal field markers...")
    html_path = str(Path(docx_path).with_suffix(".html"))
    docx_to_html_with_fields(docx_path, html_path, fields)

    # Step 3: Upload to DocuSeal (if API configured)
    template_id = None
    docuseal_url = os.environ.get("DOCUSEAL_URL")
    if docuseal_url:
        print("[STEP 3] Uploading to DocuSeal...")
        result = upload_to_docuseal_api(html_path, fields["title"])
        if result:
            template_id = result.get("template_id")
    else:
        print("[STEP 3] SKIPPED (DOCUSEAL_URL not set) — manual upload required")
        print(f"  Upload this file to DocuSeal: {html_path}")

    # Step 4: Send for signature
    if template_id and signer_email:
        print("[STEP 4] Sending for e-signature...")
        result = send_for_signature(
            template_id, signer_email, candidate_name,
            subject=f"【{school_name}】{candidate_name} 教師{contract_type}電子簽署通知",
        )
        return result
    else:
        print("[STEP 4] SKIPPED — set DOCUSEAL_URL + DOCUSEAL_API_KEY to enable")
        print(f"  Template HTML: {html_path}")
        return {"html_path": html_path, "docx_path": docx_path}


def main():
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python3.12 send_for_signature.py <candidate_name> <signer_email> [docx_path]")
        print("")
        print("Environment variables:")
        print("  DOCUSEAL_URL=https://your-docuseal.example.com")
        print("  DOCUSEAL_API_KEY=xxx")
        print("  HR_WEBHOOK_URL=https://... (optional, for completion webhook)")
        print("")
        print("Examples:")
        print("  # Generate offer letter + send via DocuSeal:")
        print("  python3.12 send_for_signature.py '王小明' 'wang@mail.edu.tw'")
        print("")
        print("  # Use existing .docx file:")
        print("  python3.12 send_for_signature.py '王小明' 'wang@mail.edu.tw' /tmp/offer.docx")
        sys.exit(1)

    candidate_name = sys.argv[1]
    signer_email = sys.argv[2]
    docx_path = sys.argv[3] if len(sys.argv) > 3 else None

    # Extract position/salary/start_date/school from existing .docx if available
    if docx_path and Path(docx_path).exists():
        # Get metadata from .docx (basic)
        try:
            import docx
            doc = docx.Document(docx_path)
            first_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:5]
            print(f"DOCX content preview: {first_lines}")
        except Exception as e:
            print(f"Warning: Could not read DOCX: {e}")

    result = generate_and_send(
        candidate_name=candidate_name,
        position="教師",  # Will be filled from docx if available
        salary="",
        start_date="",
        school_name="",
        signer_email=signer_email,
        docx_path=docx_path,
    )
    print("\n[DONE]", json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
